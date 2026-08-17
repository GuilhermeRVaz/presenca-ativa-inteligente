import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls

MD_CONTENT = """# Relatório Oficial de Operação — Busca Ativa (11/08/2026)
## Faltas do Dia 11/08/2026 e Follow-up | EE Décia — Presença Ativa Inteligente (PAI)

---

> ℹ️ **CONTEXTO OPERACIONAL DA CAMPANHA DE 11/08/2026**
>
> Este relatório apresenta a **auditoria consolidada e oficial da operação de Busca Ativa e Follow-up realizada em 11/08/2026**, referente às **ausências registradas no próprio dia 11/08/2026** (ID Campanha Principal: `04890cfd-d4aa-4d6b-9e12-cf87439f13df` | ID Follow-up: `48a4e8bd-f09e-4117-be32-18605a796e06`).
>
> A operação cobriu um universo de **25 estudantes faltosos** de 7 turmas (6ºA ao 9ºA). Foi aplicada a **Orquestração de Disparos v3.0 com Pacing de Distribuição Log-Normal** (delays de 40s a 180s, além de rajadas e pausas de comportamento humano natural).
>
> Ao todo, **12 famílias interagiram e responderam**, atingindo **52,2% de taxa de resposta útil** (12 respostas sobre 23 mensagens efetivamente entregues). Destaca-se a recuperação de justificativa via disparo de **Follow-up para a aluna Lara Lima de Souza (8ºB)**, identificando um caso crítico de saúde mental.

---

## 📊 1. Resumo Executivo e Métricas da Campanha (Faltas 11/08/2026)

- **Data das Ausências:** 11/08/2026 (Terça-feira)
- **ID Campanha Principal:** `04890cfd-d4aa-4d6b-9e12-cf87439f13df` (Busca Ativa 11/08)
- **ID Campanha Follow-up:** `48a4e8bd-f09e-4117-be32-18605a796e06` (Follow-up 11/08)
- **Execução dos Disparos Primários:** 11/08/2026 — Entre 09h22 e 10h36
- **Execução do Follow-up:** 11/08/2026 — Entre 13h14 e 13h51
- **Motor de Disparo:** Orquestrador v3.0 (Log-Normal Pacing Engine + Micro-pausas)
- **Status Operacional:** Campanha finalizada com consolidado completo auditado no banco de dados

| Métrica Operacional | Quantidade | Percentual / Detalhe |
|---|---|---|
| **Total de Faltosos Identificados no Dia (11/08)** | **25** | 100% — universo alvo da campanha |
| **Mensagens Enviadas com Sucesso (Disparo Primário)** | **23** | 92,0% de entrega efetiva no WhatsApp |
| **Falhas de Envio API (Número Inexistente no WA)** | **2** | 8,0% — Samuel Saraiva (6ºA) e Luis Miguel (9ºA) |
| **Mensagens Enviadas na 2ª Onda (Follow-up)** | **11** | 100% de sucesso nos contatos secundários elegíveis |
| **Respostas Coletadas (Famílias que Interagiram)** | **12** | **52,2%** de engajamento (12 / 23 entregues) |
| **Justificativas Válidas / Motivos Identificados** | **11** | 47,8% com motivo claro registrado |
| **Respostas de Confirmação Geral (Sem Motivo Detalhado)** | **1** | 4,3% — Sofia Filassi Prado (7ºB) |
| **Alunos Silenciosos (Sem Resposta após 2 Ondas)** | **11** | 47,8% dos contatados (necessitam contato direto) |

---

## 📋 2. Detalhamento por Turma — Todos os 25 Faltosos de 11/08

### 6º ANO 6A (2 alunos)

| Nº | Aluno(a) | RA | Protocolo | Status | Resposta da Família | Categoria |
|---|---|---|---|---|---|---|
| 1 | LUCAS SALVADOR SANTANA RIBEIRO | 121119695 | P-180687 | ✅ Respondeu | *"Boa tarde P-180687 hoje ele não foi porque ontem a tarde ele machucou o pulso"* | Saúde / Lesão (Pulso) |
| 2 | SAMUEL HONORIO DE ALMEIDA SARAIVA | 115939931 | — | 🔴 Falha API | *"Erro 400: número 5514997005093 não possui WhatsApp"* | Cadastro Incorreto |

### 6º ANO 6B (1 aluno)

| Nº | Aluno(a) | RA | Protocolo | Status | Resposta da Família | Categoria |
|---|---|---|---|---|---|---|
| 3 | PEDRO HENRIQUE PEIXOTO DE DEUS | 116654637 | P-C7436E | ✅ Respondeu | *"P-C7436E ele machucou o pé jogando bola , por isso preciso levar ao médico pra fazer rx"* | Saúde / Trauma (Pé / Raio-X) |

### 7º ANO 7A (5 alunos)

| Nº | Aluno(a) | RA | Protocolo | Status | Resposta da Família | Categoria |
|---|---|---|---|---|---|---|
| 4 | ISABELLA DE LIMA CARVALHO TOMÉ | 113930336 | P-F1B624 | ✅ Respondeu | *"P-F1B624 Isabella Carvalho Tome Houve uma intercorrência pela manhã não sendo possível levá-la."* | Imprevisto / Logística |
| 5 | ISADORA MEIRA DA COSTA SILVA | 113794622 | P-B97BAD | ✅ Respondeu | *"Renite atacada"* | Saúde (Rinite) |
| 6 | LIVIA MARTINS CRUZ | 114982434 | P-9D2A93 | 📤 Sem Resposta | — | Pendente (Primário + Follow-up) |
| 7 | SARAH GABRIELLA SILVA MANOEL | 114658396 | P-0C7C3B | 📤 Sem Resposta | — | Pendente (Primário + Follow-up) |
| 8 | LUIZ FABIANO LEONCIO DE GODOY | 114658338 | P-F6E029 | 📤 Sem Resposta | — | Pendente (Primário + Follow-up) |

### 7º ANO 7B (6 alunos)

| Nº | Aluno(a) | RA | Protocolo | Status | Resposta da Família | Categoria |
|---|---|---|---|---|---|---|
| 9 | ANA JULIA LIMA DE SOUZA | 120213171 | P-7964BB | ✅ Respondeu | *"Ana julia Lima de Souza Cólica menstrual"* | Saúde (Cólica) |
| 10 | DENYKA RIHANNA AMORIM LEME | 114886343 | P-DE6046 | ✅ Respondeu | *"P-DE6046 Denyka Rihanna amorim leme esta com febre"* | Saúde (Febre) |
| 11 | NICOLAS FERNANDO GONÇALVES | 114503219 | P-436D51 | 📤 Sem Resposta | — | Pendente (Primário + Follow-up) |
| 12 | NICOLLY LUZ MARTINS | 115912977 | P-1E62A1 | ✅ Respondeu | *"Bom dia A nicolly não estava passando bem"* | Saúde (Mal-estar) |
| 13 | SOFIA FILASSI PRADO | 115867875 | P-F33D05 | 🟢 Respondeu | *"Bom-dia"* | Interação sem detalhe |
| 14 | HELOISA APARECIDA VIANA DA SILVA | 115315522 | P-C9A551 | 📤 Sem Resposta | — | Pendente (Primário + Follow-up) |

### 8º ANO 8A (3 alunos)

| Nº | Aluno(a) | RA | Protocolo | Status | Resposta da Família | Categoria |
|---|---|---|---|---|---|---|
| 15 | GIOVANA DUARTE RODRIGUES | 112756430 | P-4BB7F0 | ✅ Respondeu | *"Hoje a Giovana acordou indisposta... Dia de passar em consulta médica."* | Saúde / Consulta |
| 16 | LUIZ ANTONIO NOBREGA NETO | 116100257 | P-52BA5D | ✅ Respondeu | *"Nao acordou bem hj"* | Saúde (Mal-estar) |
| 17 | RAFAEL HENRIQUE CORREA | 112786408 | P-8814BA | ✅ Respondeu | *"Ele está com dor de cabeça / A manha vou ai"* | Saúde (Dor de cabeça) |

### 8º ANO 8B (6 alunos)

| Nº | Aluno(a) | RA | Protocolo | Status | Resposta da Família | Categoria |
|---|---|---|---|---|---|---|
| 18 | LARA LIMA DE SOUZA | 114722206 | P-A2C3E8 | 🚨 Respondeu | *"P-A2C3E8 A Lara está com problemas psicológicos | Acho que ela está com início de depressão, estou vendo uma psicóloga pra ela"* | **SAÚDE MENTAL (Depressão)** |
| 19 | SOPHIA GABRIELE ROSA DE ARRUDA | 114900397 | P-0E58B9 | 📤 Sem Resposta | — | Pendente (Primário + Follow-up) |
| 20 | SOPHIA RAYANE DE CARVALHO | 116083865 | P-9D1B18 | 📤 Sem Resposta | — | Pendente (Primário + Follow-up) |
| 21 | KAUE PEREIRA DE OLIVEIRA | 114063618 | P-DEBD90 | 📤 Sem Resposta | — | Pendente (Primário + Follow-up) |
| 22 | CECILIA MARTINS CRUZ | 114982423 | P-734212 | 📤 Sem Resposta | — | Pendente (Primário + Follow-up) |
| 23 | DAVI ROBERTO DA SILVA | 120753622 | P-11C300 | 📤 Sem Resposta | — | Pendente (Sem 2º contato) |

### 9º ANO 9A (2 alunos)

| Nº | Aluno(a) | RA | Protocolo | Status | Resposta da Família | Categoria |
|---|---|---|---|---|---|---|
| 24 | LUIS MIGUEL PAIVA BERNAVA | 112807059 | — | 🔴 Falha API | *"Erro 400: número 5514998608582 não possui WhatsApp"* | Cadastro Incorreto |
| 25 | ISIS QUINTANILHA SAMPAIO | 111571004 | P-AF547C | 📤 Sem Resposta | — | Pendente (Primário + Follow-up) |

---

## ✅ 3. Consolidado de Justificativas Coletadas (12 Alunos — 11/08/2026)

| # | Aluno(a) | Turma | Justificativa Literal / Resposta | Categoria | Encaminhamento |
|---|---|---|---|---|---|
| 1 | **LARA LIMA DE SOUZA** | 8ºB | *"P-A2C3E8 A Lara está com problemas psicológicos | Acho que ela está com início de depressão, estou vendo uma psicóloga pra ela"* | **Saúde Mental / Depressão** | 🚨 **URGENTE:** Acompanhamento Psicossocial / Tutoria |
| 2 | PEDRO HENRIQUE PEIXOTO DE DEUS | 6ºB | *"P-C7436E ele machucou o pé jogando bola , por isso preciso levar ao médico pra fazer rx"* | Saúde / Trauma (Pé / RX) | Aguardar comprovante/atestado médico |
| 3 | LUCAS SALVADOR SANTANA RIBEIRO | 6ºA | *"Boa tarde P-180687 hoje ele não foi porque ontem a tarde ele machucou o pulso"* | Saúde / Lesão (Pulso) | Registrar justificativa física |
| 4 | DENYKA RIHANNA AMORIM LEME | 7ºB | *"P-DE6046 Denyka Rihanna amorim leme esta com febre"* | Saúde (Febre) | Lançar falta por motivo de saúde |
| 5 | ANA JULIA LIMA DE SOUZA | 7ºB | *"Ana julia Lima de Souza Cólica menstrual"* | Saúde (Cólica) | Lançar falta por motivo de saúde |
| 6 | ISADORA MEIRA DA COSTA SILVA | 7ºA | *"Renite atacada"* | Saúde (Rinite) | Lançar falta por motivo de saúde |
| 7 | GIOVANA DUARTE RODRIGUES | 8ºA | *"Hoje a Giovana acordou indisposta... Dia de passar em consulta médica."* | Saúde / Consulta | Aguardar atestado/comprovante de consulta |
| 8 | RAFAEL HENRIQUE CORREA | 8ºA | *"Ele está com dor de cabeça / A manha vou ai"* | Saúde (Dor de cabeça) | Aluno esteve na escola e foi retirado |
| 9 | NICOLLY LUZ MARTINS | 7ºB | *"Bom dia A nicolly não estava passando bem"* | Saúde (Mal-estar) | Lançar falta por motivo de saúde |
| 10 | LUIZ ANTONIO NOBREGA NETO | 8ºA | *"Nao acordou bem hj"* | Saúde (Mal-estar) | Lançar falta por motivo de saúde |
| 11 | ISABELLA DE LIMA CARVALHO TOMÉ | 7ºA | *"P-F1B624 Isabella Carvalho Tome Houve uma intercorrência pela manhã não sendo possível levá-la."* | Imprevisto / Logística | Justificativa de rotina familiar |
| 12 | SOFIA FILASSI PRADO | 7ºB | *"Bom-dia"* | Interação Geral | Confirmação de recebimento |

---

## 🛑 4. Alunos Silenciosos (11 Casos Sem Resposta — 11/08)

| # | Aluno(a) | Turma | RA | Status WhatsApp | Encaminhamento |
|---|---|---|---|---|---|
| 1 | LIVIA MARTINS CRUZ | 7ºA | 114982434 | 📤 Entregue (2 Ondas) | Contato telefônico direto da secretaria |
| 2 | SARAH GABRIELLA SILVA MANOEL | 7ºA | 114658396 | 📤 Entregue (2 Ondas) | Contato telefônico com responsável |
| 3 | LUIZ FABIANO LEONCIO DE GODOY | 7ºA | 114658338 | 📤 Entregue (2 Ondas) | Ligação de acompanhamento da secretaria |
| 4 | NICOLAS FERNANDO GONÇALVES | 7ºB | 114503219 | 📤 Entregue (2 Ondas) | Contato telefônico com mãe/pai |
| 5 | HELOISA APARECIDA VIANA DA SILVA | 7ºB | 115315522 | 📤 Entregue (2 Ondas) | Verificar histórico de frequência |
| 6 | SOPHIA GABRIELE ROSA DE ARRUDA | 8ºB | 114900397 | 📤 Entregue (2 Ondas) | Verificar processo de transferência anterior |
| 7 | SOPHIA RAYANE DE CARVALHO | 8ºB | 116083865 | 📤 Entregue (2 Ondas) | Contato telefônico com responsável |
| 8 | KAUE PEREIRA DE OLIVEIRA | 8ºB | 114063618 | 📤 Entregue (2 Ondas) | Contato com vô/mãe |
| 9 | CECILIA MARTINS CRUZ | 8ºB | 114982423 | 📤 Entregue (2 Ondas) | Contato telefônico com pai/mãe |
| 10 | DAVI ROBERTO DA SILVA | 8ºB | 120753622 | 📤 Entregue (1 Onda) | Sem 2º contato cadastrado — ligar |
| 11 | ISIS QUINTANILHA SAMPAIO | 9ºA | 111571004 | 📤 Entregue (2 Ondas) | Contato direto via secretaria |

---

## 🚨 5. Casos de Atenção Especial (11/08/2026)

| Prioridade | Aluno(a) | Turma | Motivo | Ação Recomendada |
|---|---|---|---|---|
| 🚨 **URGENTE** | **LARA LIMA DE SOUZA** | 8ºB | Mãe relatou início de depressão e acompanhamento psicológico | Encaminhar imediatamente para Mediação Escolar e Tutoria para suporte psicossocial |
| 🏥 **ALTA** | **PEDRO HENRIQUE PEIXOTO DE DEUS** | 6ºB | Machucou o pé jogando bola e foi fazer raio-X no hospital | Registrar justificativa médica e acompanhar retorno às aulas |
| 🏥 **ALTA** | **LUCAS SALVADOR SANTANA RIBEIRO** | 6ºA | Lesão no pulso ocorrida na tarde anterior | Registrar falta por motivo de saúde |
| ⚠️ **MÉDIA** | **SAMUEL HONORIO DE ALMEIDA SARAIVA** | 6ºA | Erro 400 de transmissão: número de telefone sem WhatsApp | Atualizar número do responsável no sistema SEDUC |
| ⚠️ **MÉDIA** | **LUIS MIGUEL PAIVA BERNAVA** | 9ºA | Erro 400 de transmissão: número de telefone sem WhatsApp | Atualizar número do responsável no sistema SEDUC |
| 🔵 **INFO** | **RAFAEL HENRIQUE CORREA** | 8ºA | Queixou-se de dor de cabeça na escola e foi retirado pelo irmão | Confirmar com a família se haverá atestado para 11/08 |
| 🔵 **INFO** | **GIOVANA DUARTE RODRIGUES** | 8ºA | Indisposição e consulta médica agendada | Solicitar comprovante de consulta médica |

---

## 📌 6. Encaminhamentos e Orientações Pedagógicas

1. **🚨 Suporte Psicossocial Urgente (Lara Lima de Souza - 8ºB):** Articular entre a coordenação pedagógica, tutoria e mediação escolar um plano de escuta e acolhimento acolhedor para a aluna, mantendo contato próximo com a mãe para apoiar no tratamento psicológico.
2. **🏥 Acompanhamento de Traumas e Atestados (Pedro Henrique - 6ºB e Lucas Salvador - 6ºA):** Lançar as faltas justificadas no diário de classe e solicitar a entrega dos atestados ou comprovantes médicos assim que os alunos retornarem.
3. **🔧 Atualização Cadastral Urgente na SEDUC:** Corrigir os números de telefone das famílias de **Samuel Honorio (6ºA)** e **Luis Miguel Paiva Bernava (9ºA)**, cujos disparos falharam por ausência de conta ativa no WhatsApp.
4. **📞 Busca Ativa Direta dos 11 Silenciosos:** Acionar a equipe de secretaria para contato telefônico direto com os 11 responsáveis que não responderam às duas ondas de disparo automatizado de 11/08.

---

*Documento gerado em: 12/08/2026 | Sistema: Presença Ativa Inteligente (PAI) | EE Professora Décia Vilela Ramos*
"""

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'<w:tblBorders {nsdecls("w")}><w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:left w:val="none"/><w:right w:val="none"/><w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:insideV w:val="none"/></w:tblBorders>')
    tblPr.append(borders)

def build_docx():
    doc = Document()
    
    # Page setup - Margins 2cm
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Base Styles
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x26, 0x26, 0x26)

    # Title Banner / Header
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("RELATÓRIO OFICIAL DE BUSCA ATIVA — 11/08/2026")
    run_title.bold = True
    run_title.font.size = Pt(20)
    run_title.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78) # Navy Blue

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(14)
    run_sub = sub_p.add_run("Escola Estadual Décia — Sistema Presença Ativa Inteligente (PAI)")
    run_sub.font.size = Pt(13)
    run_sub.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    # Callout Box - CONTEXTO OPERACIONAL
    callout_tbl = doc.add_table(rows=1, cols=1)
    callout_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    callout_cell = callout_tbl.cell(0, 0)
    set_cell_background(callout_cell, "F2F5F8") # Light blue-gray box
    set_cell_margins(callout_cell, top=140, bottom=140, left=200, right=200)
    
    tcPr = callout_cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="1F4E78"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
    tcPr.append(borders)

    cp = callout_cell.paragraphs[0]
    cp.paragraph_format.space_before = Pt(2)
    cp.paragraph_format.space_after = Pt(2)
    run_alert_title = cp.add_run("ℹ️ CONTEXTO OPERACIONAL DA CAMPANHA DE 11/08/2026:\n")
    run_alert_title.bold = True
    run_alert_title.font.size = Pt(11)
    run_alert_title.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

    run_alert_body = cp.add_run(
        "Este relatório apresenta a auditoria consolidada e oficial da operação de Busca Ativa e Follow-up realizada em 11/08/2026, "
        "referente às ausências registradas no próprio dia 11/08/2026 (IDs: 04890cfd-d4aa-4d6b-9e12-cf87439f13df e 48a4e8bd-f09e-4117-be32-18605a796e06).\n"
        "A operação mobilizou 25 alunos faltosos com disparo primário e 11 disparos de reiteração (follow-up). Foi aplicada a "
        "Orquestração v3.0 com pacing de Distribuição Log-Normal e comportamento humano natural. A taxa de resposta útil alcançou 52,2%."
    )
    run_alert_body.font.size = Pt(10.5)
    run_alert_body.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Heading 1: Resumo Executivo
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    r_h1 = h1.add_run("1. Resumo Executivo e Métricas da Campanha (Faltas 11/08/2026)")
    r_h1.bold = True
    r_h1.font.size = Pt(14)
    r_h1.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

    # Meta Bullets
    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_after = Pt(10)
    p_meta.paragraph_format.line_spacing = 1.15
    
    r = p_meta.add_run("• Data das Ausências: ")
    r.bold = True
    p_meta.add_run("11/08/2026 (Terça-feira)\n")
    
    r = p_meta.add_run("• ID Campanha Principal (1ª Onda): ")
    r.bold = True
    p_meta.add_run("04890cfd-d4aa-4d6b-9e12-cf87439f13df\n")

    r = p_meta.add_run("• ID Campanha Follow-up (2ª Onda): ")
    r.bold = True
    p_meta.add_run("48a4e8bd-f09e-4117-be32-18605a796e06\n")

    r = p_meta.add_run("• Horário de Execução: ")
    r.bold = True
    p_meta.add_run("Primário: 09h22 – 10h36 | Follow-up: 13h14 – 13h51\n")

    r = p_meta.add_run("• Status Operacional: ")
    r.bold = True
    p_meta.add_run("Concluído (Auditado via banco Supabase e Evolution API)")

    # Metrics Table
    m_table = doc.add_table(rows=9, cols=3)
    m_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(m_table, color="D9D9D9", sz="4")

    headers = ["Métrica Operacional", "Quantidade", "Percentual / Detalhe"]
    hdr_cells = m_table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1F4E78")
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.size = Pt(10.5)

    metrics_data = [
        ("Total de Faltosos Identificados (Alvo)", "25", "100,0% — universo alvo da campanha"),
        ("Mensagens Enviadas no Primário (WhatsApp)", "23", "92,0% de taxa de transmissão bem-sucedida"),
        ("Falhas de Transmissão API (Telefone Inexistente)", "2", "8,0% — Samuel Saraiva (6ºA) e Luis Miguel (9ºA)"),
        ("Mensagens Enviadas no Follow-up (2º Contato)", "11", "100,0% de entrega nos contatos elegíveis"),
        ("Respostas Coletadas (Famílias que Interagiram)", "12", "52,2% de engajamento útil (12 / 23 entregues)"),
        ("Justificativas Válidas Coletadas", "11", "47,8% com motivo claro registrado"),
        ("Respostas Gerais / Sem Motivo Detalhado", "1", "4,3% — Sofia Filassi Prado (7ºB)"),
        ("Alunos Silenciosos (Sem Resposta)", "11", "47,8% dos contatados (para busca ativa presencial)")
    ]

    for row_idx, (m_title, m_val, m_desc) in enumerate(metrics_data, start=1):
        row_cells = m_table.rows[row_idx].cells
        bg_color = "F2F5F8" if row_idx % 2 == 1 else "FFFFFF"
            
        row_cells[0].text = m_title
        row_cells[1].text = m_val
        row_cells[2].text = m_desc

        for col_idx in range(3):
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=100, bottom=100, left=150, right=150)
            p = row_cells[col_idx].paragraphs[0]
            if col_idx == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Heading 2: Detalhamento por Turma
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(6)
    r_h2 = h2.add_run("2. Detalhamento por Turma — Todos os 25 Alunos Faltosos")
    r_h2.bold = True
    r_h2.font.size = Pt(14)
    r_h2.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

    students_table = doc.add_table(rows=26, cols=6)
    students_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(students_table, color="D9D9D9", sz="4")

    s_headers = ["Aluno(a)", "Turma", "Protocolo", "Status", "Resposta da Família", "Categoria"]
    s_hdr_cells = students_table.rows[0].cells
    for i, title in enumerate(s_headers):
        s_hdr_cells[i].text = title
        set_cell_background(s_hdr_cells[i], "1F4E78")
        set_cell_margins(s_hdr_cells[i], top=120, bottom=120, left=100, right=100)
        p = s_hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.size = Pt(9.5)

    all_students_data = [
        ("LUCAS SALVADOR SANTANA RIBEIRO", "6º A", "P-180687", "✅ Respondeu", "machucou o pulso ontem a tarde", "Saúde / Lesão"),
        ("SAMUEL HONORIO DE ALMEIDA SARAIVA", "6º A", "—", "🔴 Falha API", "Número 5514997005093 inexistente no WA", "Cadastro Incorreto"),
        ("PEDRO HENRIQUE PEIXOTO DE DEUS", "6º B", "P-C7436E", "✅ Respondeu", "machucou o pé jogando bola, fazer rx", "Saúde / Trauma"),
        ("ISABELLA DE LIMA CARVALHO TOMÉ", "7º A", "P-F1B624", "✅ Respondeu", "Intercorrência pela manhã não sendo possível levá-la", "Imprevisto Famil."),
        ("ISADORA MEIRA DA COSTA SILVA", "7º A", "P-B97BAD", "✅ Respondeu", "Renite atacada", "Saúde (Rinite)"),
        ("LIVIA MARTINS CRUZ", "7º A", "P-9D2A93", "📤 Sem Resp.", "—", "Pendente"),
        ("SARAH GABRIELLA SILVA MANOEL", "7º A", "P-0C7C3B", "📤 Sem Resp.", "—", "Pendente"),
        ("LUIZ FABIANO LEONCIO DE GODOY", "7º A", "P-F6E029", "📤 Sem Resp.", "—", "Pendente"),
        ("ANA JULIA LIMA DE SOUZA", "7º B", "P-7964BB", "✅ Respondeu", "Cólica menstrual", "Saúde (Cólica)"),
        ("DENYKA RIHANNA AMORIM LEME", "7º B", "P-DE6046", "✅ Respondeu", "Denyka esta com febre", "Saúde (Febre)"),
        ("NICOLAS FERNANDO GONÇALVES", "7º B", "P-436D51", "📤 Sem Resp.", "—", "Pendente"),
        ("NICOLLY LUZ MARTINS", "7º B", "P-1E62A1", "✅ Respondeu", "A nicolly não estava passando bem", "Saúde (Mal-estar)"),
        ("SOFIA FILASSI PRADO", "7º B", "P-F33D05", "🟢 Respondeu", "Bom-dia", "Interação Geral"),
        ("HELOISA APARECIDA VIANA DA SILVA", "7º B", "P-C9A551", "📤 Sem Resp.", "—", "Pendente"),
        ("GIOVANA DUARTE RODRIGUES", "8º A", "P-4BB7F0", "✅ Respondeu", "Giovana acordou indisposta. Consulta médica.", "Saúde / Consulta"),
        ("LUIZ ANTONIO NOBREGA NETO", "8º A", "P-52BA5D", "✅ Respondeu", "Nao acordou bem hj", "Saúde (Mal-estar)"),
        ("RAFAEL HENRIQUE CORREA", "8º A", "P-8814BA", "✅ Respondeu", "Dor de cabeça / Irmão foi buscar na escola", "Saúde (Cabeça)"),
        ("LARA LIMA DE SOUZA", "8º B", "P-A2C3E8", "🚨 URGENTE", "Problemas psicológicos / início de depressão", "SAÚDE MENTAL"),
        ("SOPHIA GABRIELE ROSA DE ARRUDA", "8º B", "P-0E58B9", "📤 Sem Resp.", "—", "Pendente"),
        ("SOPHIA RAYANE DE CARVALHO", "8º B", "P-9D1B18", "📤 Sem Resp.", "—", "Pendente"),
        ("KAUE PEREIRA DE OLIVEIRA", "8º B", "P-DEBD90", "📤 Sem Resp.", "—", "Pendente"),
        ("CECILIA MARTINS CRUZ", "8º B", "P-734212", "📤 Sem Resp.", "—", "Pendente"),
        ("DAVI ROBERTO DA SILVA", "8º B", "P-11C300", "📤 Sem Resp.", "—", "Pendente"),
        ("LUIS MIGUEL PAIVA BERNAVA", "9º A", "—", "🔴 Falha API", "Número 5514998608582 inexistente no WA", "Cadastro Incorreto"),
        ("ISIS QUINTANILHA SAMPAIO", "9º A", "P-AF547C", "📤 Sem Resp.", "—", "Pendente")
    ]

    for row_idx, (aluno, turma, proto, st, msg, cat) in enumerate(all_students_data, start=1):
        row_cells = students_table.rows[row_idx].cells
        if "URGENTE" in st or "SAÚDE MENTAL" in cat:
            bg_color = "FCE4D6" # Soft orange / alert
        elif "Falha" in st:
            bg_color = "FFF2CC" # Soft yellow
        elif row_idx % 2 == 1:
            bg_color = "F2F5F8"
        else:
            bg_color = "FFFFFF"

        row_cells[0].text = aluno
        row_cells[1].text = turma
        row_cells[2].text = proto
        row_cells[3].text = st
        row_cells[4].text = msg
        row_cells[5].text = cat

        for col_idx in range(6):
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=70, bottom=70, left=80, right=80)
            p = row_cells[col_idx].paragraphs[0]
            p.runs[0].font.size = Pt(8.5)
            if col_idx in (1, 2, 3):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Heading 3: Consolidado de Justificativas
    h3 = doc.add_paragraph()
    h3.paragraph_format.space_before = Pt(14)
    h3.paragraph_format.space_after = Pt(6)
    r_h3 = h3.add_run("3. Consolidado de Justificativas Coletadas (12 Alunos)")
    r_h3.bold = True
    r_h3.font.size = Pt(14)
    r_h3.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

    j_table = doc.add_table(rows=13, cols=5)
    j_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(j_table, color="D9D9D9", sz="4")

    j_headers = ["Aluno(a)", "Turma", "Justificativa Literal / Resposta", "Categoria", "Encaminhamento"]
    j_hdr_cells = j_table.rows[0].cells
    for i, title in enumerate(j_headers):
        j_hdr_cells[i].text = title
        set_cell_background(j_hdr_cells[i], "1F4E78")
        set_cell_margins(j_hdr_cells[i], top=120, bottom=120, left=100, right=100)
        p = j_hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.size = Pt(9.5)

    justifications = [
        ("LARA LIMA DE SOUZA", "8º B", "P-A2C3E8 A Lara está com problemas psicológicos | Acho que ela está com início de depressão, estou vendo uma psicóloga pra ela", "Saúde Mental / Depressão", "🚨 URGENTE: Mediação Escolar / Tutoria / Acolhimento"),
        ("PEDRO HENRIQUE PEIXOTO DE DEUS", "6º B", "P-C7436E ele machucou o pé jogando bola , por isso preciso levar ao médico pra fazer rx", "Saúde / Trauma", "Aguardar comprovante/atestado médico"),
        ("LUCAS SALVADOR SANTANA RIBEIRO", "6º A", "Boa tarde P-180687 hoje ele não foi porque ontem a tarde ele machucou o pulso", "Saúde / Lesão Pulso", "Lançar falta por motivo de saúde"),
        ("DENYKA RIHANNA AMORIM LEME", "7º B", "P-DE6046 Denyka Rihanna amorim leme esta com febre", "Saúde (Febre)", "Lançar falta por motivo de saúde"),
        ("ANA JULIA LIMA DE SOUZA", "7º B", "Ana julia Lima de Souza Cólica menstrual", "Saúde (Cólica)", "Lançar falta por motivo de saúde"),
        ("ISADORA MEIRA DA COSTA SILVA", "7º A", "Renite atacada", "Saúde (Rinite)", "Lançar falta por motivo de saúde"),
        ("GIOVANA DUARTE RODRIGUES", "8º A", "Hoje a Giovana acordou indisposta... Dia de passar em consulta médica.", "Saúde / Consulta", "Solicitar comprovante de consulta médica"),
        ("RAFAEL HENRIQUE CORREA", "8º A", "Ele está com dor de cabeça / A manha vou ai", "Saúde (Dor de cabeça)", "Aluno esteve na escola e foi retirado pelo irmão"),
        ("NICOLLY LUZ MARTINS", "7º B", "Bom dia A nicolly não estava passando bem", "Saúde (Mal-estar)", "Lançar falta por motivo de saúde"),
        ("LUIZ ANTONIO NOBREGA NETO", "8º A", "Nao acordou bem hj", "Saúde (Mal-estar)", "Lançar falta por motivo de saúde"),
        ("ISABELLA DE LIMA CARVALHO TOMÉ", "7º A", "P-F1B624 Isabella Carvalho Tome Houve uma intercorrência pela manhã...", "Imprevisto Logística", "Justificativa familiar aceita"),
        ("SOFIA FILASSI PRADO", "7º B", "Bom-dia", "Interação Geral", "Confirmação de recebimento")
    ]

    for row_idx, (aluno, turma, msg, cat, enc) in enumerate(justifications, start=1):
        row_cells = j_table.rows[row_idx].cells
        if "Depressão" in cat or "URGENTE" in enc:
            bg_color = "FCE4D6"
        elif row_idx % 2 == 1:
            bg_color = "F2F5F8"
        else:
            bg_color = "FFFFFF"

        row_cells[0].text = aluno
        row_cells[1].text = turma
        row_cells[2].text = msg
        row_cells[3].text = cat
        row_cells[4].text = enc

        for col_idx in range(5):
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=100, right=100)
            p = row_cells[col_idx].paragraphs[0]
            p.runs[0].font.size = Pt(8.5)
            if col_idx == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Heading 4: Alunos Silenciosos
    h4 = doc.add_paragraph()
    h4.paragraph_format.space_before = Pt(14)
    h4.paragraph_format.space_after = Pt(6)
    r_h4 = h4.add_run("4. Alunos Silenciosos (11 Casos Sem Resposta)")
    r_h4.bold = True
    r_h4.font.size = Pt(14)
    r_h4.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

    nr_table = doc.add_table(rows=12, cols=4)
    nr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(nr_table, color="D9D9D9", sz="4")

    nr_headers = ["Aluno(a)", "Turma", "Status WhatsApp", "Encaminhamento Recomendado"]
    nr_hdr_cells = nr_table.rows[0].cells
    for i, title in enumerate(nr_headers):
        nr_hdr_cells[i].text = title
        set_cell_background(nr_hdr_cells[i], "595959")
        set_cell_margins(nr_hdr_cells[i], top=120, bottom=120, left=120, right=120)
        p = nr_hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.size = Pt(9.5)

    no_responses = [
        ("LIVIA MARTINS CRUZ", "7º ANO 7A", "📤 Entregue (Primário + Follow-up)", "Contato telefônico direto da secretaria"),
        ("SARAH GABRIELLA SILVA MANOEL", "7º ANO 7A", "📤 Entregue (Primário + Follow-up)", "Contato telefônico com responsável"),
        ("LUIZ FABIANO LEONCIO DE GODOY", "7º ANO 7A", "📤 Entregue (Primário + Follow-up)", "Ligação de acompanhamento da secretaria"),
        ("NICOLAS FERNANDO GONÇALVES", "7º ANO 7B", "📤 Entregue (Primário + Follow-up)", "Contato telefônico com mãe/pai"),
        ("HELOISA APARECIDA VIANA DA SILVA", "7º ANO 7B", "📤 Entregue (Primário + Follow-up)", "Verificar histórico de frequência"),
        ("SOPHIA GABRIELE ROSA DE ARRUDA", "8º ANO 8B", "📤 Entregue (Primário + Follow-up)", "Verificar processo de transferência"),
        ("SOPHIA RAYANE DE CARVALHO", "8º ANO 8B", "📤 Entregue (Primário + Follow-up)", "Contato telefônico com responsável"),
        ("KAUE PEREIRA DE OLIVEIRA", "8º ANO 8B", "📤 Entregue (Primário + Follow-up)", "Contato telefônico com vô/mãe"),
        ("CECILIA MARTINS CRUZ", "8º ANO 8B", "📤 Entregue (Primário + Follow-up)", "Contato telefônico com pai/mãe"),
        ("DAVI ROBERTO DA SILVA", "8º ANO 8B", "📤 Entregue (Primário)", "Sem 2º contato — realizar ligação"),
        ("ISIS QUINTANILHA SAMPAIO", "9º ANO 9A", "📤 Entregue (Primário + Follow-up)", "Contato direto via secretaria")
    ]

    for row_idx, (aluno, turma, status_wa, rec) in enumerate(no_responses, start=1):
        row_cells = nr_table.rows[row_idx].cells
        bg_color = "F9F9F9" if row_idx % 2 == 1 else "FFFFFF"

        row_cells[0].text = aluno
        row_cells[1].text = turma
        row_cells[2].text = status_wa
        row_cells[3].text = rec

        for col_idx in range(4):
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=100, right=100)
            p = row_cells[col_idx].paragraphs[0]
            p.runs[0].font.size = Pt(8.5)
            if col_idx in (1, 2):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Heading 5: Casos de Atenção Especial
    h5 = doc.add_paragraph()
    h5.paragraph_format.space_before = Pt(14)
    h5.paragraph_format.space_after = Pt(6)
    r_h5 = h5.add_run("5. Casos de Atenção Especial (11/08/2026)")
    r_h5.bold = True
    r_h5.font.size = Pt(14)
    r_h5.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

    att_table = doc.add_table(rows=8, cols=5)
    att_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(att_table, color="D9D9D9", sz="4")

    att_headers = ["Prioridade", "Aluno(a)", "Turma", "Motivo / Relato", "Ação Recomendada"]
    att_hdr_cells = att_table.rows[0].cells
    for i, title in enumerate(att_headers):
        att_hdr_cells[i].text = title
        set_cell_background(att_hdr_cells[i], "1F4E78")
        set_cell_margins(att_hdr_cells[i], top=120, bottom=120, left=100, right=100)
        p = att_hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.size = Pt(9.5)

    att_cases = [
        ("🚨 URGENTE", "LARA LIMA DE SOUZA", "8º B", "Mãe relatou início de depressão e acompanhamento psicológico", "Encaminhar para Tutoria, Mediação Escolar e Acolhimento Psicossocial"),
        ("🏥 ALTA", "PEDRO HENRIQUE PEIXOTO DE DEUS", "6º B", "Machucou o pé jogando bola / foi fazer raio-X no hospital", "Registrar falta justificada e aguardar atestado/comprovante médico"),
        ("🏥 ALTA", "LUCAS SALVADOR SANTANA RIBEIRO", "6º A", "Lesão no pulso ocorrida na tarde anterior", "Registrar falta por motivo de saúde"),
        ("⚠️ MÉDIA", "SAMUEL HONORIO DE ALMEIDA SARAIVA", "6º A", "Erro 400 de transmissão: telefone sem conta de WhatsApp", "Atualizar cadastro do responsável no sistema SEDUC"),
        ("⚠️ MÉDIA", "LUIS MIGUEL PAIVA BERNAVA", "9º A", "Erro 400 de transmissão: telefone sem conta de WhatsApp", "Atualizar cadastro do responsável no sistema SEDUC"),
        ("🔵 INFO", "RAFAEL HENRIQUE CORREA", "8º A", "Sentiu dor de cabeça na escola e foi retirado pelo irmão", "Acompanhar retorno e confirmar atestado"),
        ("🔵 INFO", "GIOVANA DUARTE RODRIGUES", "8º A", "Indisposição / Consulta médica agendada", "Solicitar comprovante de consulta médica")
    ]

    for row_idx, (prio, aluno, turma, motivo, acao) in enumerate(att_cases, start=1):
        row_cells = att_table.rows[row_idx].cells
        if "URGENTE" in prio:
            bg_color = "FCE4D6"
        elif "ALTA" in prio:
            bg_color = "FDF7F2"
        else:
            bg_color = "FFFFFF"

        row_cells[0].text = prio
        row_cells[1].text = aluno
        row_cells[2].text = turma
        row_cells[3].text = motivo
        row_cells[4].text = acao

        for col_idx in range(5):
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=100, right=100)
            p = row_cells[col_idx].paragraphs[0]
            p.runs[0].font.size = Pt(8.5)
            if col_idx in (0, 2):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Heading 6: Encaminhamentos Pedagógicos
    h6 = doc.add_paragraph()
    h6.paragraph_format.space_before = Pt(14)
    h6.paragraph_format.space_after = Pt(6)
    r_h6 = h6.add_run("6. Encaminhamentos e Orientações Pedagógicas")
    r_h6.bold = True
    r_h6.font.size = Pt(14)
    r_h6.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

    recs = [
        ("🚨 Suporte Psicossocial Urgente: ", "Articular entre a coordenação pedagógica, tutoria e mediação escolar um plano de escuta e acolhimento para a aluna Lara Lima de Souza (8ºB), mantendo contato próximo com a mãe."),
        ("🏥 Acompanhamento de Traumas e Saúde: ", "Registrar no diário de classe as faltas por motivo de saúde/trauma de Pedro Henrique (6ºB), Lucas Salvador (6ºA), Denyka Rihanna (7ºB), Ana Júlia (7ºB) e Isadora Meira (7ºA)."),
        ("🔧 Atualização Cadastral SEDUC: ", "Regularizar os contatos de WhatsApp das famílias de Samuel Honorio (6ºA) e Luis Miguel Paiva Bernava (9ºA)."),
        ("📞 Busca Ativa dos 11 Silenciosos: ", "Priorizar ligação telefônica direta da secretaria para os 11 responsáveis que não responderam às duas ondas de disparo automatizado.")
    ]

    for title, desc in recs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.2)
        r_t = p.add_run(f"• {title}")
        r_t.bold = True
        r_t.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
        r_d = p.add_run(desc)
        r_d.font.size = Pt(10.5)

    # Footer
    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_foot = p_foot.add_run("EE Décia — Presença Ativa Inteligente (PAI) | Documento Gerado em 12/08/2026")
    r_foot.font.size = Pt(9)
    r_foot.font.italic = True
    r_foot.font.color.rgb = RGBColor(0x8C, 0x8C, 0x8C)

    # Output paths
    out_dir = Path("relatorios")
    out_dir.mkdir(exist_ok=True)
    
    file_docx1 = out_dir / "RELATORIO_OFICIAL_BUSCA_ATIVA_11_08_FALTAS_DO_DIA.docx"
    file_docx2 = out_dir / "RELATORIO_OFICIAL_BUSCA_ATIVA_DIA_11_08.docx"
    file_md = out_dir / "RELATORIO_OFICIAL_BUSCA_ATIVA_DIA_11_08.md"

    # Write Markdown
    with open(file_md, "w", encoding="utf-8") as f:
        f.write(MD_CONTENT)
    print(f"[OK] MD GERADO: {file_md.resolve()}")

    # Save DOCX (both standard names for user convenience)
    doc.save(str(file_docx1))
    print(f"[OK] DOCX GERADO: {file_docx1.resolve()}")
    
    doc.save(str(file_docx2))
    print(f"[OK] DOCX GERADO: {file_docx2.resolve()}")

if __name__ == "__main__":
    build_docx()
