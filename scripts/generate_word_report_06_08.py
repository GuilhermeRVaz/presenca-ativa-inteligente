import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'<w:tblBorders {nsdecls("w")}><w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:left w:val="none"/><w:right w:val="none"/><w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:insideV w:val="none"/></w:tblBorders>')
    tblPr.append(borders)

def build_docx():
    doc = Document()
    
    # Page setup - Margins 2cm
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Styles
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x26, 0x26, 0x26)

    # Title Banner / Header
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("RELATÓRIO OFICIAL DE BUSCA ATIVA — 06/08/2026")
    run_title.bold = True
    run_title.font.size = Pt(20)
    run_title.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78) # Navy Blue

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(14)
    run_sub = sub_p.add_run("Escola Estadual Décia — Sistema Presença Ativa Inteligente (PAI)")
    run_sub.font.size = Pt(13)
    run_sub.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    # Callout Box - NOTA OPERACIONAL (SEM FOLLOW-UP)
    callout_tbl = doc.add_table(rows=1, cols=1)
    callout_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    callout_cell = callout_tbl.cell(0, 0)
    set_cell_background(callout_cell, "FFF2CC") # Soft yellow alert background
    set_cell_margins(callout_cell, top=140, bottom=140, left=200, right=200)
    
    # Left border highlight in amber
    tcPr = callout_cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="C55A11"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
    tcPr.append(borders)

    cp = callout_cell.paragraphs[0]
    cp.paragraph_format.space_before = Pt(2)
    cp.paragraph_format.space_after = Pt(2)
    run_alert_title = cp.add_run("⚠️ NOTA OPERACIONAL IMPORTANTE — ENVIO DE FOLLOW-UP:\n")
    run_alert_title.bold = True
    run_alert_title.font.size = Pt(11)
    run_alert_title.font.color.rgb = RGBColor(0xC5, 0x5A, 0x11) # Amber dark

    run_alert_body = cp.add_run(
        "Nesta campanha referente às faltas do dia 06/08/2026, foi realizada exclusivamente a primeira onda de disparo principal. "
        "NÃO FOI REALIZADO O ENVIO DE FOLLOW-UP (segunda tentativa de contato) para os alunos pendentes. "
        "Todos os dados e taxas de resposta deste relatório refletem unicamente a abordagem inicial efetuada pelo sistema."
    )
    run_alert_body.font.size = Pt(10.5)
    run_alert_body.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Heading 1: Resumo Executivo
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    r_h1 = h1.add_run("1. Resumo Executivo e Métricas da Campanha")
    r_h1.bold = True
    r_h1.font.size = Pt(14)
    r_h1.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

    # Info Bullet Points / Meta
    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_after = Pt(10)
    p_meta.paragraph_format.line_spacing = 1.15
    
    r = p_meta.add_run("• Data da Campanha / Ausências: ")
    r.bold = True
    p_meta.add_run("06/08/2026\n")
    
    r = p_meta.add_run("• ID da Campanha no Supabase: ")
    r.bold = True
    p_meta.add_run("8f9bab32-4d45-4a4b-badf-f20c88b08027\n")

    r = p_meta.add_run("• Status Operacional: ")
    r.bold = True
    p_meta.add_run("Concluído (Sem disparo de reiteração / follow-up)")

    # Metrics Table
    m_table = doc.add_table(rows=6, cols=3)
    m_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(m_table, color="D9D9D9", sz="4")

    headers = ["Métrica Operacional", "Quantidade", "Percentual / Detalhe"]
    hdr_cells = m_table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1F4E78")
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.size = Pt(10.5)

    metrics_data = [
        ("Total de Alunos Selecionados (Alvo)", "28", "100,0% das faltas registradas"),
        ("Mensagens Enviadas com Sucesso", "24", "85,7% das linhas processadas"),
        ("Falhas de Transmissão (Sem Tel. Cadastrado)", "4", "14,3% necessitam atualização na SEDUC"),
        ("Respostas / Interações Coletadas", "14", "58,3% de engajamento das famílias"),
        ("Justificativas Válidas Coletadas", "12", "50,0% com motivo de falta identificado")
    ]

    for row_idx, (m_title, m_val, m_desc) in enumerate(metrics_data, start=1):
        row_cells = m_table.rows[row_idx].cells
        if row_idx % 2 == 1:
            bg_color = "F2F5F8"
        else:
            bg_color = "FFFFFF"
            
        row_cells[0].text = m_title
        row_cells[1].text = m_val
        row_cells[2].text = m_desc

        for col_idx in range(3):
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=100, bottom=100, left=150, right=150)
            p = row_cells[col_idx].paragraphs[0]
            if col_idx == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Heading 2: Detalhamento de Justificativas
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(6)
    r_h2 = h2.add_run("2. Detalhamento de Justificativas e Retornos Coletados (14 Famílias)")
    r_h2.bold = True
    r_h2.font.size = Pt(14)
    r_h2.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

    j_table = doc.add_table(rows=15, cols=6)
    j_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(j_table, color="D9D9D9", sz="4")

    j_headers = ["Aluno(a)", "Turma", "Protocolo", "Justificativa / Mensagem da Família", "Categoria", "Status Protocolo"]
    j_hdr_cells = j_table.rows[0].cells
    for i, title in enumerate(j_headers):
        j_hdr_cells[i].text = title
        set_cell_background(j_hdr_cells[i], "1F4E78")
        set_cell_margins(j_hdr_cells[i], top=120, bottom=120, left=120, right=120)
        p = j_hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.size = Pt(9.5)

    justifications = [
        ("DIOGO COSSO SILVA", "9º A", "P-2D5B06", "Virose. Ele está com virose.", "Saúde", "🔑 Confirmado"),
        ("MARIA EDUARDA DE OLIVEIRA", "7º B", "P-CE07B8", "Consulta Médica. Amanhã ela leva declaração.", "Saúde / Médico", "🔑 Confirmado"),
        ("RAFAEL YATSUDA AZEVEDO", "8º A", "P-703BB5", "Motivo particular.", "Particular", "🔑 Confirmado"),
        ("SÔNIA MARIA CARUZO FIGUEIREDO", "6º B", "P-6B7E66", "O carro da família apresentou problemas mecânicos e não deu para levá-la para a escola.", "Logística Familiar", "🔑 Confirmado"),
        ("LARA LIMA DE SOUZA", "8º B", "P-B7C544", "A Lara não acordou bem hoje, acordou vomitando por isso não mandei para a escola mais amanhã ela está de volta.", "Saúde", "🔑 Confirmado"),
        ("ALESSANDRA KAROLINE PEREIRA DE SOUZA", "8º B", "P-220B2A", "Boa tarde, emocional não muito legal para ir não ter ido ontem e hoje. Motivos da crise de ansiedade.", "Saúde Mental / Emocional", "🟢 Identificado"),
        ("GUILHERME DE OLIVEIRA MARIN", "8º A", "P-9D806A", "Guilherme foi até 11:00.", "Saída Antecipada", "🟢 Identificado"),
        ("LARA BEATRIZ DO NASCIMENTO SOUZA", "8º B", "P-B1FCD2", "Boa tarde. Lara Beatriz do Nascimento Souza. Ela perdeu o horário.", "Logística / Atraso", "🟢 Identificado"),
        ("ANDRÉ SANTANA DE BRITO", "9º A", "P-F296C7", "Boa tarde. Ele teve dentista.", "Odontológico", "🟢 Identificado"),
        ("GABRIELLY BENTO ALVES", "9º A", "P-C6D44A", "Oi boa tarde. Ela estava com dor na barriga.", "Saúde", "🟢 Identificado"),
        ("ANA JULIA LIMA DE SOUZA", "7º B", "P-DD0ACC", "Ana Júlia Lima de Souza. Dor de cabeça.", "Saúde", "🟢 Identificado"),
        ("MAYTE COELHO DIAS", "8º B", "P-29C3EB", "Tive que ir a Bauru no período da manhã com retorno a tarde após as 16h, optamos pela falta para não ficar sozinha em casa.", "Viagem / Logística", "🟢 Identificado"),
        ("LUIZ ANTONIO NOBREGA NETO", "8º A", "P-5B51C4", "Perdeu hora.", "Atraso", "❓ Sem Motivo Claro"),
        ("ENZO GABRIEL FERREIRA DE OLIVEIRA", "8º A", "P-5B0E81", "Mensagem vazia / confirmação de recebimento.", "Acompanhamento", "❓ Sem Motivo Claro")
    ]

    for row_idx, (aluno, turma, proto, msg, cat, status) in enumerate(justifications, start=1):
        row_cells = j_table.rows[row_idx].cells
        bg_color = "F2F5F8" if row_idx % 2 == 1 else "FFFFFF"

        row_cells[0].text = aluno
        row_cells[1].text = turma
        row_cells[2].text = proto
        row_cells[3].text = msg
        row_cells[4].text = cat
        row_cells[5].text = status

        for col_idx in range(6):
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=100, right=100)
            p = row_cells[col_idx].paragraphs[0]
            p.runs[0].font.size = Pt(9)
            if col_idx in (1, 2, 5):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Heading 3: Alunos Sem Resposta
    h3 = doc.add_paragraph()
    h3.paragraph_format.space_before = Pt(14)
    h3.paragraph_format.space_after = Pt(6)
    r_h3 = h3.add_run("3. Alunos Sem Resposta (10 Casos Pendentes)")
    r_h3.bold = True
    r_h3.font.size = Pt(14)
    r_h3.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

    p_no_resp_note = doc.add_paragraph()
    p_no_resp_note.paragraph_format.space_after = Pt(8)
    r_note = p_no_resp_note.add_run(
        "Como não houve disparo de follow-up, os 10 alunos abaixo permaneceram em status de aguardo de retorno após o primeiro envio do WhatsApp:"
    )
    r_note.font.italic = True
    r_note.font.size = Pt(10)

    nr_table = doc.add_table(rows=11, cols=4)
    nr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(nr_table, color="D9D9D9", sz="4")

    nr_headers = ["Aluno(a)", "Turma", "Status WhatsApp", "Encaminhamento Recomendado"]
    nr_hdr_cells = nr_table.rows[0].cells
    for i, title in enumerate(nr_headers):
        nr_hdr_cells[i].text = title
        set_cell_background(nr_hdr_cells[i], "595959") # Dark Grey Header
        set_cell_margins(nr_hdr_cells[i], top=120, bottom=120, left=120, right=120)
        p = nr_hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.size = Pt(9.5)

    no_responses = [
        ("EMILLY BEATRIZ GIMENEZ DA SILVA", "6º ANO 6A", "📤 Entregue", "Ligação / Contato direto da secretaria"),
        ("RAYSSA VITORIA DA SILVA RUBIN", "6º ANO 6A", "📤 Entregue", "Verificar histórico com família (Jaú)"),
        ("YURI GONÇALO DA SILVA", "7º ANO 7A", "📤 Entregue", "Acompanhamento presencial"),
        ("LUCAS GABRIEL ZINI DEFENDI", "6º ANO 6B", "📤 Entregue", "Contato direto com o responsável"),
        ("ARTHUR GABRIEL DIAS PIOVESAN", "9º ANO 9A", "📤 Entregue", "Verificação com tutoria/coordenação"),
        ("LARA HONORIO AMORIM", "7º ANO 7B", "📤 Entregue", "Acompanhar retorno na próxima aula"),
        ("MATHEUS HENRIQUE DOS REIS SANTANA", "8º ANO 8B", "📤 Entregue", "Contato telefônico complementar"),
        ("JULIA DUTRA MENDROTE", "7º ANO 7B", "📤 Entregue", "Acompanhamento familiar"),
        ("ANDRERY FERNANDES DE OLIVEIRA", "8º ANO 8B", "📤 Entregue", "Verificar atestado médico de gripe"),
        ("ISIS QUINTANILHA SAMPAIO", "9º ANO 9A", "📤 Entregue", "Acompanhar pedido de transferência")
    ]

    for row_idx, (aluno, turma, status_wa, rec) in enumerate(no_responses, start=1):
        row_cells = nr_table.rows[row_idx].cells
        bg_color = "F9F9F9" if row_idx % 2 == 1 else "FFFFFF"

        row_cells[0].text = aluno
        row_cells[1].text = turma
        row_cells[2].text = status_wa
        row_cells[3].text = rec

        for col_idx in range(4):
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=100, right=100)
            p = row_cells[col_idx].paragraphs[0]
            p.runs[0].font.size = Pt(9)
            if col_idx in (1, 2):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Heading 4: Falhas de Transmissão
    h4 = doc.add_paragraph()
    h4.paragraph_format.space_before = Pt(14)
    h4.paragraph_format.space_after = Pt(6)
    r_h4 = h4.add_run("4. Falhas de Transmissão (4 Alunos Sem Telefone Valido)")
    r_h4.bold = True
    r_h4.font.size = Pt(14)
    r_h4.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

    p_fail_intro = doc.add_paragraph()
    p_fail_intro.paragraph_format.space_after = Pt(6)
    p_fail_intro.add_run(
        "Os 4 alunos abaixo não puderam receber a mensagem de busca ativa devido à ausência ou inconsistência do número de telefone cadastrado no sistema:"
    )

    f_table = doc.add_table(rows=5, cols=3)
    f_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(f_table, color="D9D9D9", sz="4")

    f_headers = ["Aluno(a)", "Turma", "Ação Necessária"]
    f_hdr_cells = f_table.rows[0].cells
    for i, title in enumerate(f_headers):
        f_hdr_cells[i].text = title
        set_cell_background(f_hdr_cells[i], "C55A11") # Amber Header
        set_cell_margins(f_hdr_cells[i], top=120, bottom=120, left=120, right=120)
        p = f_hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.size = Pt(9.5)

    failures = [
        ("ANA JULYA BARBOSA DE OLIVEIRA", "8º ANO 8A", "Atualizar telefone no cadastro SEDUC"),
        ("BRAYAN DUTRA MENDROTE", "8º ANO 8B", "Atualizar telefone no cadastro SEDUC"),
        ("MARIA EDUARDA PAES ALVES", "8º ANO 8A", "Atualizar telefone no cadastro SEDUC"),
        ("PEDRO MIGUEL MIRANDA LONGO", "8º ANO 8B", "Atualizar telefone no cadastro SEDUC")
    ]

    for row_idx, (aluno, turma, acao) in enumerate(failures, start=1):
        row_cells = f_table.rows[row_idx].cells
        bg_color = "FDF7F2" if row_idx % 2 == 1 else "FFFFFF"

        row_cells[0].text = aluno
        row_cells[1].text = turma
        row_cells[2].text = acao

        for col_idx in range(3):
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=100, right=100)
            p = row_cells[col_idx].paragraphs[0]
            p.runs[0].font.size = Pt(9)
            if col_idx == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Heading 5: Recomendações
    h5 = doc.add_paragraph()
    h5.paragraph_format.space_before = Pt(14)
    h5.paragraph_format.space_after = Pt(6)
    r_h5 = h5.add_run("5. Encaminhamentos e Recomendações Pedagogicas")
    r_h5.bold = True
    r_h5.font.size = Pt(14)
    r_h5.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

    recs = [
        ("Registro de Atestados de Saúde: ", "Lançar no diário de classe a justificativa médica e virose informada pelas famílias de Diogo Cosso Silva (9ºA), Maria Eduarda de Oliveira (7ºB) e Lara Lima de Souza (8ºB)."),
        ("Acompanhamento Emocional: ", "Encaminhar para a mediação escolar / tutoria a aluna Alessandra Karoline Pereira de Souza (8ºB), cuja mãe informou faltas decorrentes de crises de ansiedade."),
        ("Acompanhamento de Atrasos: ", "Orientar as famílias de Lara Beatriz do Nascimento (8ºB) e Luiz Antonio Nóbrega Neto (8ºA) sobre o cumprimento dos horários de entrada."),
        ("Atualização Cadastral Urgente: ", "Regularizar os contatos de WhatsApp dos 4 alunos cujas mensagens não puderam ser entregues.")
    ]

    for title, desc in recs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.2)
        r_t = p.add_run(f"• {title}")
        r_t.bold = True
        r_t.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
        r_d = p.add_run(desc)
        r_d.font.size = Pt(10.5)

    # Footer
    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_foot = p_foot.add_run("EE Décia — Presença Ativa Inteligente (PAI) | Documento Gerado em 07/08/2026")
    r_foot.font.size = Pt(9)
    r_foot.font.italic = True
    r_foot.font.color.rgb = RGBColor(0x8C, 0x8C, 0x8C)

    # Output path
    out_dir = Path("relatorios")
    out_dir.mkdir(exist_ok=True)
    out_file = out_dir / "Relatorio_Oficial_Busca_Ativa_06_08_2026.docx"
    doc.save(str(out_file))
    print(f"[OK] RELATORIO WORD GERADO COM SUCESSO: {out_file.resolve()}")

if __name__ == "__main__":
    build_docx()
